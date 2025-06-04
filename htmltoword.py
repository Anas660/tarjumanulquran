import os
import glob
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt

def convert_html_to_word(html_path, word_path):
    # Read the HTML file
    with open(html_path, 'r', encoding='utf-8') as file:
        html_content = file.read()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create a Word document
    doc = Document()
    
    # Set RTL direction for the entire document (for Urdu text)
    for section in doc.sections:
        section.page_width = section.page_width  # This forces page setup to be applied
    
    # Find the title
    title_element = soup.find(['h1', 'h2', 'title', 'h3'], class_=['entry-title', 'post-title', 'title'])
    if title_element:
        title = title_element.get_text(strip=True)
        heading = doc.add_heading(title, level=1)
        # Set RTL alignment for the title
        for run in heading.runs:
            run.font.rtl = True
            run.font.size = Pt(16)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Debug: Print what we found before processing
    print(f"Processing file: {html_path}")
    
    # First, find all accordion sections with their titles
    accordion_sections = []
    
    # Method 1: Look for accordion buttons/titles and their associated content
    accordion_buttons = soup.find_all(['button', 'a'], class_='accordion-button')
    for button in accordion_buttons:
        title_text = button.get_text(strip=True)
        # Find the corresponding accordion content div
        # Usually it's the next sibling or in a data-target reference
        content_id = button.get('data-target') or button.get('href')
        content_div = None
        
        if content_id and content_id.startswith('#'):
            # If there's a reference to an ID, find that element
            content_div = soup.find(id=content_id[1:])
        else:
            # Otherwise try finding the next accordion-desc sibling
            parent = button.parent
            if parent:
                content_div = parent.find_next('div', class_='accordion-desc')
        
        if content_div:
            accordion_sections.append((title_text, content_div))
    
    # Method 2: Look for structured accordion with title in header followed by content
    accordion_headers = soup.find_all(['h2', 'h3', 'h4', 'div'], class_=['accordion-header', 'card-header'])
    for header in accordion_headers:
        title_text = header.get_text(strip=True)
        # Content is usually the next sibling or next div
        content_div = header.find_next('div', class_='accordion-desc')
        if content_div:
            accordion_sections.append((title_text, content_div))
    
    # Method 3: Direct approach - find accordion-desc divs and look for preceding title
    accordion_descs = soup.find_all('div', class_='accordion-desc')
    for desc in accordion_descs:
        # Try to find a title element before this accordion-desc
        prev_elem = desc.find_previous(['h2', 'h3', 'h4', 'h5', 'button', 'a', 'span'])
        title_text = prev_elem.get_text(strip=True) if prev_elem else f"Section {accordion_descs.index(desc) + 1}"
        accordion_sections.append((title_text, desc))
    
    print(f"Found {len(accordion_sections)} accordion sections")
    
    # If no accordion sections found, use the fallback method (similar to previous script)
    if not accordion_sections:
        all_paragraphs = []
        # Find paragraphs in entry-content
        entry_content = soup.find('div', class_='entry-content')
        if entry_content:
            p_tags = entry_content.find_all('p', recursive=True)
            all_paragraphs.extend(p_tags)
        
        # If still no paragraphs, look throughout the document
        if not all_paragraphs:
            all_paragraphs = soup.find_all('p')
        
        # Process unique paragraphs
        unique_paragraphs = []
        seen = set()
        for p in all_paragraphs:
            p_text = p.get_text().strip()
            if p_text and p_text not in seen and len(p_text) > 10:
                unique_paragraphs.append(p)
                seen.add(p_text)
        
        print(f"No accordion sections found. Processing {len(unique_paragraphs)} direct paragraphs instead.")
        
        for p_tag in unique_paragraphs:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            text = p_tag.get_text(strip=True)
            if not text:
                continue
                
            run = p.add_run(text)
            run.font.rtl = True
            
            # Apply formatting (color, bold, italic)
            if 'style' in p_tag.attrs and 'color:' in p_tag['style']:
                if '#ff0000' in p_tag['style']:
                    run.font.color.rgb = RGBColor(255, 0, 0)
                elif '#008000' in p_tag['style']:
                    run.font.color.rgb = RGBColor(0, 128, 0)
            
            if p_tag.find(['strong', 'b']):
                run.bold = True
            if p_tag.find(['em', 'i']):
                run.italic = True
    else:
        # Process accordion sections - title followed by content
        for title, content_div in accordion_sections:
            # Add accordion title as heading
            heading = doc.add_heading(title, level=2)
            for run in heading.runs:
                run.font.rtl = True
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Process content paragraphs
            p_tags = content_div.find_all('p', recursive=True)
            unique_paragraphs = []
            seen = set()
            
            for p in p_tags:
                p_text = p.get_text().strip()
                if p_text and p_text not in seen and len(p_text) > 10:
                    unique_paragraphs.append(p)
                    seen.add(p_text)
            
            print(f"Found {len(unique_paragraphs)} paragraphs in section '{title[:20]}...'")
            
            for p_tag in unique_paragraphs:
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
                text = p_tag.get_text(strip=True)
                if not text:
                    continue
                    
                run = p.add_run(text)
                run.font.rtl = True
                
                # Apply formatting (color, bold, italic)
                if 'style' in p_tag.attrs and 'color:' in p_tag['style']:
                    if '#ff0000' in p_tag['style']:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    elif '#008000' in p_tag['style']:
                        run.font.color.rgb = RGBColor(0, 128, 0)
                
                # Check for spans with color
                for span in p_tag.find_all('span', style=True):
                    if 'color:' in span['style']:
                        if '#ff0000' in span['style']:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                        elif '#008000' in span['style']:
                            run.font.color.rgb = RGBColor(0, 128, 0)
                
                if p_tag.find(['strong', 'b']):
                    run.bold = True
                if p_tag.find(['em', 'i']):
                    run.italic = True
    
    # Save the Word document
    doc.save(word_path)
    print(f"Converted {html_path} to {word_path}")

def main():
    # Path to the folder containing HTML files
    html_folder = "maududi_books_html"
    
    # Create output folder for Word documents if it doesn't exist
    word_folder = "maududi_books_word"
    if not os.path.exists(word_folder):
        os.makedirs(word_folder)
    
    # Get all HTML files in the folder
    html_files = glob.glob(os.path.join(html_folder, "*.html"))
    
    print(f"Found {len(html_files)} HTML files to convert...")
    
    # Convert each HTML file to Word
    for html_file in html_files:
        # Get the filename without extension
        base_name = os.path.basename(html_file)
        file_name = os.path.splitext(base_name)[0]
        
        # Create output Word file path
        word_file = os.path.join(word_folder, file_name + ".docx")
        
        # Convert HTML to Word
        try:
            convert_html_to_word(html_file, word_file)
        except Exception as e:
            print(f"Error converting {html_file}: {e}")
    
    print(f"Successfully converted HTML files to Word documents.")
    print(f"Word documents are saved in the '{word_folder}' folder.")

if __name__ == "__main__":
    main()