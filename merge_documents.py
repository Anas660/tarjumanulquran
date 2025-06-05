import os
import glob
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.enum.section import WD_SECTION_START

# Input and output directories
INPUT_DIR = "rasailomasail_word"
OUTPUT_DIR = "rasailomasail_merged"

def create_output_dir():
    """Create output directory if it doesn't exist"""
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

def copy_element_formatting(source_paragraph, target_paragraph):
    """Copy formatting from source paragraph to target paragraph"""
    # Copy alignment
    target_paragraph.alignment = source_paragraph.alignment
    
    # Copy paragraph style if it exists and is valid in the target document
    if source_paragraph.style:
        try:
            target_paragraph.style = source_paragraph.style.name
        except:
            pass
    
    # Copy paragraph formatting
    target_paragraph.paragraph_format.left_indent = source_paragraph.paragraph_format.left_indent
    target_paragraph.paragraph_format.right_indent = source_paragraph.paragraph_format.right_indent
    target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
    target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
    target_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing

def merge_volume_documents(volume_num):
    """Merge all documents in a volume into a single document"""
    volume_dir = os.path.join(INPUT_DIR, f"volume_{volume_num:02d}")
    output_file = os.path.join(OUTPUT_DIR, f"volume_{volume_num:02d}_merged.docx")
    
    print(f"\nMerging documents in Volume {volume_num}...")
    
    if not os.path.exists(volume_dir):
        print(f"Volume directory {volume_dir} not found. Skipping.")
        return False
    
    # Get all Word files in the volume directory
    word_files = glob.glob(os.path.join(volume_dir, '*.docx'))
    
    if not word_files:
        print(f"No Word documents found in {volume_dir}. Skipping.")
        return False
    
    print(f"Found {len(word_files)} documents to merge")
    
    # Create a new document for the merged output
    merged_doc = Document()
    
    # Configure document for RTL (Urdu)
    for section in merged_doc.sections:
        section.page_width = section.page_width  # This forces page setup to be applied
    
    # Add volume title
    title = f"مجموعہ رسائل و مسائل - جلد {volume_num}"  # "Collection of Rasail-o-Masail - Volume X" in Urdu
    heading = merged_doc.add_heading(title, level=0)
    for run in heading.runs:
        run.font.rtl = True
        run.font.size = Pt(20)
        run.bold = True
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Add table of contents heading
    toc_heading = merged_doc.add_heading("فہرست مضامین", level=1)  # "Table of Contents" in Urdu
    for run in toc_heading.runs:
        run.font.rtl = True
        run.font.size = Pt(16)
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Create table of contents
    toc = []
    
    # First pass: collect all article titles for table of contents
    print("Building table of contents...")
    for i, doc_path in enumerate(word_files, 1):
        try:
            doc = Document(doc_path)
            
            # Get the first heading as the title
            title = f"مضمون {i}"  # Default title: "Article X" in Urdu
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    title = para.text
                    break
            
            # Add to TOC list
            toc.append((title, i))
            
        except Exception as e:
            print(f"Error reading document {doc_path}: {e}")
            toc.append((f"مضمون {i}", i))  # Add default entry in case of error
    
    # Add table of contents entries
    for title, article_num in toc:
        toc_para = merged_doc.add_paragraph()
        toc_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add article number
        run = toc_para.add_run(f"{article_num}. ")
        run.font.rtl = True
        
        # Add article title
        run = toc_para.add_run(title)
        run.font.rtl = True
    
    # Add a page break after TOC
    merged_doc.add_page_break()
    
    # Second pass: merge documents
    print("Merging documents...")
    for i, doc_path in enumerate(word_files, 1):
        try:
            # Add article number and title as a heading
            article_title = toc[i-1][0]
            article_heading = merged_doc.add_heading(f"{i}. {article_title}", level=1)
            for run in article_heading.runs:
                run.font.rtl = True
                run.font.size = Pt(16)
            article_heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Open source document
            doc = Document(doc_path)
            
            # Skip the first heading as we've already added it
            skip_first_heading = True
            
            # Copy all paragraphs from source doc to merged doc
            for para in doc.paragraphs:
                # Skip the first heading (title) as we already added it with the article number
                if skip_first_heading and para.style.name.startswith('Heading'):
                    skip_first_heading = False
                    continue
                
                # Copy paragraph with its formatting
                p = merged_doc.add_paragraph()
                copy_element_formatting(para, p)
                
                # Copy all runs with their formatting
                for run in para.runs:
                    new_run = p.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.rtl = True  # Ensure RTL direction
                    
                    # Copy font properties
                    if run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
                    if run.font.size:
                        new_run.font.size = run.font.size
            
            # Add a page break between articles
            if i < len(word_files):
                merged_doc.add_page_break()
            
            print(f"Added article {i}/{len(word_files)}: {os.path.basename(doc_path)}")
            
        except Exception as e:
            print(f"Error processing document {doc_path}: {e}")
            # Add error note in the merged document
            p = merged_doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run = p.add_run(f"Error including document: {os.path.basename(doc_path)}")
            run.font.rtl = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red text for error
            
            # Add a page break before continuing to the next document
            if i < len(word_files):
                merged_doc.add_page_break()
    
    # Save the merged document
    merged_doc.save(output_file)
    print(f"Successfully created merged document: {output_file}")
    return True

def main():
    """Main function to merge Word documents by volume"""
    print("Starting to merge Word documents by volume...")
    
    # Create output directory
    create_output_dir()
    
    # Process each volume
    successful = 0
    for volume_num in range(1, 6):
        if merge_volume_documents(volume_num):
            successful += 1
    
    print(f"\nMerging complete! {successful}/5 volume documents created successfully.")
    print(f"Merged documents are saved in the '{OUTPUT_DIR}' folder.")

if __name__ == "__main__":
    main()