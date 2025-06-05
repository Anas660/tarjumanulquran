import os
import glob
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt
from urllib.parse import unquote

# Input and output directories
INPUT_DIR = "rasailomasail_articles"
OUTPUT_DIR = "rasailomasail_word"

def create_output_dirs():
    """Create output directory for Volume 5 Word documents"""
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
    
    # Create only volume 5 directory
    vol_dir = os.path.join(OUTPUT_DIR, "volume_05")
    if not os.path.exists(vol_dir):
        os.makedirs(vol_dir)

def get_readable_filename(html_filename):
    """Convert the URL-encoded filename to a more readable one"""
    # Remove .html extension
    base_name = os.path.splitext(html_filename)[0]
    
    # URL decode the filename (convert %xx sequences to characters)
    decoded = unquote(base_name)
    
    # Replace any remaining problematic chars with underscores
    safe_name = ''.join(c if c.isalnum() or c in '-_ ' else '_' for c in decoded)
    
    return safe_name + '.docx'

def convert_html_to_word(html_path, word_path):
    """Convert HTML article to Word document with proper formatting"""
    # Read the HTML file
    with open(html_path, 'r', encoding='utf-8') as file:
        html_content = file.read()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create a Word document
    doc = Document()
    
    # Set RTL direction for the entire document (for Urdu text)
    for section in doc.sections:
        section.page_width = section.page_width  # This forces page setup to be applied
    
    # Find the article title - try different patterns
    title_element = None
    possible_title_selectors = [
        ('h1', {'class_': 'entry-title'}),
        ('h1', {'class_': 'post-title'}),
        ('h1', {'class_': 'article-title'}),
        ('h2', {'class_': 'entry-title'}),
        ('h1', {}),  # Any h1
        ('header', {})  # Header element
    ]
    
    for selector, attrs in possible_title_selectors:
        if attrs:
            title_element = soup.find(selector, **attrs)
        else:
            title_element = soup.find(selector)
        
        if title_element:
            break
    
    # Extract title text
    title = ""
    if title_element:
        title = title_element.get_text(strip=True)
    else:
        # Fallback: Use filename as title
        title = os.path.splitext(os.path.basename(html_path))[0]
        title = unquote(title)  # URL decode
    
    # Add title to document
    if title:
        heading = doc.add_heading(title, level=1)
        # Set RTL alignment for the title
        for run in heading.runs:
            run.font.rtl = True
            run.font.size = Pt(16)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Find the main content area
    content_area = None
    content_selectors = [
        ('div', {'class_': 'entry-content'}),
        ('div', {'class_': 'post-content'}),
        ('div', {'class_': 'article-content'}),
        ('div', {'class_': 'content'}),
        ('article', {}),
        ('main', {})
    ]
    
    for selector, attrs in content_selectors:
        content_area = soup.find(selector, **attrs)
        if content_area:
            break
    
    # Fallback to body if no content area found
    if not content_area:
        content_area = soup.body
        
        # Remove unwanted elements from body
        for unwanted in content_area.find_all(['nav', 'header', 'footer', 'aside', 
                                              'script', 'style', 'meta', 'link', 
                                              'form', 'iframe', 'ins']):
            unwanted.decompose()
    
    # Process content - first find all paragraphs and headings
    content_elements = []
    
    if content_area:
        content_elements = content_area.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 
                                                'blockquote', 'ul', 'ol', 'div'], recursive=True)
    
    # Process each element
    current_paragraph = None
    
    for element in content_elements:
        # Skip empty elements and navigation
        if not element.get_text(strip=True) or element.find_parent(['nav', 'footer', 'header']):
            continue
            
        # Process headings
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Skip if it's the main title we already added
            if element == title_element:
                continue
                
            # Get heading level
            level = int(element.name[1])
            heading_text = element.get_text(strip=True)
            
            # Add heading
            heading = doc.add_heading(heading_text, level=min(level+1, 9))
            for run in heading.runs:
                run.font.rtl = True
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
        # Process paragraphs
        elif element.name == 'p' or (element.name == 'div' and not element.find(['p', 'div'], recursive=False)):
            text = element.get_text(strip=True)
            if text:
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
                # Check for direct formatting on the paragraph
                run = p.add_run(text)
                run.font.rtl = True
                
                # Apply basic formatting
                if element.find(['strong', 'b']):
                    run.bold = True
                if element.find(['em', 'i']):
                    run.italic = True
                
                # Check for color styling
                for span in element.find_all('span', style=True):
                    style = span.get('style', '')
                    if 'color: #ff0000' in style or 'color:#ff0000' in style:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    elif 'color: #008000' in style or 'color:#008000' in style:
                        run.font.color.rgb = RGBColor(0, 128, 0)
        
        # Process lists
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
                text = li.get_text(strip=True)
                run = p.add_run(text)
                run.font.rtl = True
        
        # Process blockquotes
        elif element.name == 'blockquote':
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p.style = 'Quote' if 'Quote' in doc.styles else 'Normal'
            
            text = element.get_text(strip=True)
            run = p.add_run(text)
            run.font.rtl = True
    
    # Save the Word document
    doc.save(word_path)
    print(f"Converted {html_path} to {word_path}")
    return True

def process_volume_5():
    """Process all HTML articles in volume 5 directory"""
    volume_num = 5
    input_dir = os.path.join(INPUT_DIR, f"volume_{volume_num:02d}")
    output_dir = os.path.join(OUTPUT_DIR, f"volume_{volume_num:02d}")
    
    print(f"\nProcessing Volume {volume_num}...")
    
    if not os.path.exists(input_dir):
        print(f"Input directory {input_dir} not found. Skipping.")
        return
    
    # Get all HTML files in the volume directory
    html_files = glob.glob(os.path.join(input_dir, '*.html'))
    
    print(f"Found {len(html_files)} HTML articles to convert")
    
    # Process each HTML file
    successful = 0
    for i, html_file in enumerate(html_files, 1):
        # Get the filename without extension
        base_name = os.path.basename(html_file)
        # Create readable docx filename
        word_filename = get_readable_filename(base_name)
        word_path = os.path.join(output_dir, word_filename)
        
        print(f"[{i}/{len(html_files)}] Converting: {base_name} → {word_filename}")
        
        # Convert HTML to Word
        if convert_html_to_word(html_file, word_path):
            successful += 1
    
    print(f"Volume {volume_num} conversion completed: {successful}/{len(html_files)} articles converted successfully")

def main():
    """Main function to convert volume 5 HTML articles to Word documents"""
    print("Starting HTML article to Word conversion for Volume 5...")
    
    # Create output directory for volume 5
    create_output_dirs()
    
    # Process only volume 5
    process_volume_5()
    
    print("\nConversion complete! Volume 5 articles have been converted to Word documents.")

if __name__ == "__main__":
    main()